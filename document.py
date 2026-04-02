"""Word 文档封装模块."""

from __future__ import annotations

import os
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from sword.section import WordSection
from sword.format import StyleFormat


class WordDocument:
    """python-docx 的封装类，简化 Word 文档操作."""

    _doc: DocxDocument

    def __init__(self, path: str | None = None) -> None:
        """
        初始化 Word 文档.

        Args:
            path: 现有 Word 文档路径. 若为 None, 则创建新文档.
        """
        self._doc: Document = Document(path) if path else Document()
        self._page_break_between_sections: bool = True
        self._has_added_section: bool = False
        self._number_counts: dict[int, int] = {
            1: 0,
            2: 0,
            3: 0,
            4: 0,
            5: 0,
            6: 0,
            7: 0,
            8: 0,
            9: 0,
        }

    def set_page_break_between_sections(self, enabled: bool) -> None:
        """设置章节间是否分页."""
        self._page_break_between_sections = enabled

    def get_number_counts(self) -> dict[int, int]:
        """获取当前编号计数状态（供外部保存延续）."""
        return self._number_counts.copy()

    def set_start_number(self, level: int, number: int) -> None:
        """
        设置章节起始编号.

        Args:
            level: 标题级别（1-9）.
            number: 起始编号值.
        """
        self._number_counts[level] = number - 1

    def add_section(
        self, title: str | None = None, title_level: int = 1
    ) -> WordSection:
        """
        添加新章节.

        Args:
            title: 章节标题（可选）.
            title_level: 标题级别（1-9，默认 1）.

        Returns:
            WordSection: 新创建的章节对象.
        """
        if self._page_break_between_sections and self._has_added_section:
            self._doc.add_page_break()
        self._has_added_section = True
        section = WordSection(self._doc, title, self._number_counts, title_level)
        return section

    def set_table_of_contents(self) -> None:
        """在文档开头插入目录（打开后需更新域或按 F9）."""
        para = (
            self._doc.paragraphs[0]
            if self._doc.paragraphs
            else self._doc.add_paragraph()
        )
        run = para.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "separate")
        fld_char3 = OxmlElement("w:fldChar")
        fld_char3.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)
        run._r.append(fld_char3)

        self._doc.add_page_break()

    def save(self, path: str, open_after_save: bool = False) -> None:
        """
        保存文档到文件.

        Args:
            path: 输出文件路径.
            open_after_save: 保存后是否打开文件.
        """
        self._doc.save(path)
        if open_after_save:
            os.startfile(path)

    @property
    def _inner(self) -> DocxDocument:
        """访问底层的 python-docx Document 对象."""
        return self._doc

    @property
    def format(self) -> StyleFormat:
        """访问文档样式格式设置."""
        return StyleFormat(self._doc)

    def __enter__(self) -> WordDocument:
        return self

    def __exit__(self, *args: Any) -> None:
        pass
