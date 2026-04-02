"""StyleFormat 模块测试."""

from __future__ import annotations

import tempfile
import os

from sword import WordDocument, StyleFormat
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


class TestStyleFormat:
    """StyleFormat 测试类."""

    def test_create_style_format(self) -> None:
        """测试创建样式格式对象."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        assert fmt is not None

    def test_set_heading_font_name(self) -> None:
        """测试设置标题字体名称."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(1, name="宋体")

    def test_set_heading_font_size(self) -> None:
        """测试设置标题字体大小."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(1, size=16)

    def test_set_heading_font_bold(self) -> None:
        """测试设置标题加粗."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(1, bold=True)

    def test_set_heading_font_color(self) -> None:
        """测试设置标题颜色."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(1, color=(255, 0, 0))

    def test_set_heading_font_all(self) -> None:
        """测试设置标题所有字体属性."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(1, name="宋体", size=16, bold=True, color=(0, 0, 255))

    def test_set_heading_font_separate_chinese_english(self) -> None:
        """测试设置标题中英文字体分开."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(1, east_asia="宋体", ascii="Times New Roman", size=16)

    def test_set_heading_font_level2(self) -> None:
        """测试设置 Heading 2 字体格式."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(2, name="黑体", size=14, bold=True)

    def test_set_heading_font_level3(self) -> None:
        """测试设置 Heading 3 字体格式."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(3, name="楷体", size=12)

    def test_set_heading_font_only_chinese(self) -> None:
        """测试只设置中文字体."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(1, east_asia="宋体")

    def test_set_heading_font_only_english(self) -> None:
        """测试只设置英文字体."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(1, ascii="Times New Roman")

    def test_set_heading_font_with_h_ansi(self) -> None:
        """测试设置 hAnsi 字体."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_font(1, east_asia="宋体", ascii="Times New Roman", h_ansi="Arial")

    def test_set_heading_paragraph_alignment(self) -> None:
        """测试设置标题段落对齐方式."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_paragraph(1, alignment="center")

    def test_set_heading_paragraph_spacing(self) -> None:
        """测试设置标题段落间距."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_paragraph(1, space_before=12, space_after=6)

    def test_set_heading_paragraph_all(self) -> None:
        """测试设置标题所有段落属性."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_heading_paragraph(1, alignment="left", space_before=12, space_after=6)

    def test_set_normal_font(self) -> None:
        """测试设置 Normal 样式字体."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_normal_font(name="宋体", size=11)

    def test_set_normal_font_separate_chinese_english(self) -> None:
        """测试设置 Normal 中英文字体分开."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.set_normal_font(east_asia="宋体", ascii="Times New Roman", size=11)

    def test_enable_outline_level(self) -> None:
        """测试启用大纲级别."""
        doc = WordDocument()
        fmt = StyleFormat(doc._inner)
        fmt.enable_outline_level("Heading 1", 0)
        fmt.enable_outline_level("Heading 2", 1)


class TestTableFormatting:
    """表格格式测试类（通过样式设置）."""

    def test_create_style(self) -> None:
        """测试创建样式."""
        doc = WordDocument()
        style = doc.format.create_style("自定义段落样式")
        assert style is not None
        assert style.name == "自定义段落样式"

    def test_create_style_character(self) -> None:
        """测试创建字符样式."""
        doc = WordDocument()
        style = doc.format.create_style("自定义字符样式", style_type="character")
        assert style is not None
        assert style.name == "自定义字符样式"

    def test_create_style_list(self) -> None:
        """测试创建列表样式."""
        doc = WordDocument()
        style = doc.format.create_style("自定义列表样式", style_type="list")
        assert style is not None
        assert style.name == "自定义列表样式"

    def test_create_style_based_on(self) -> None:
        """测试创建基于现有样式的样式."""
        doc = WordDocument()
        style = doc.format.create_style("自定义样式", based_on="Normal")
        assert style is not None
        assert style.base_style.name == "Normal"

    def test_create_table_style(self) -> None:
        """测试创建表格样式."""
        doc = WordDocument()
        style = doc.format.create_table_style("自定义表格样式")
        assert style is not None
        assert style.name == "自定义表格样式"

    def test_create_table_style_based_on(self) -> None:
        """测试创建基于现有样式的表格样式."""
        doc = WordDocument()
        style = doc.format.create_table_style("自定义样式", based_on="Table Grid")
        assert style is not None
        assert style.base_style.name == "Table Grid"

    def test_create_table_style_exists(self) -> None:
        """测试创建已存在的表格样式返回现有样式."""
        doc = WordDocument()
        style1 = doc.format.create_table_style("测试样式")
        style2 = doc.format.create_table_style("测试样式")
        assert style1.name == style2.name

    def test_set_table_font(self) -> None:
        """测试设置表格样式字体."""
        doc = WordDocument()
        doc.format.set_table_font("Table Grid", name="宋体", size=12, bold=True)

    def test_set_table_shading(self) -> None:
        """测试设置表格样式底纹."""
        doc = WordDocument()
        doc.format.set_table_shading("Table Grid", fill="FFFF00")

    def _get_border_elem(self, doc: WordDocument, style_name: str, border_name: str):
        """获取表格样式指定边框的 XML 元素."""
        from docx.oxml.ns import qn
        style = doc.format._get_style(style_name, WD_STYLE_TYPE.TABLE)
        tblPr = style._element.find(qn("w:tblPr"))
        tblBorders = tblPr.find(qn("w:tblBorders"))
        return tblBorders.find(qn(f"w:{border_name}"))

    def test_set_table_borders(self) -> None:
        """测试设置表格样式边框."""
        doc = WordDocument()
        doc.format.set_table_borders(
            "Table Grid",
            top="single",
            bottom="single",
            left="single",
            right="single",
            inside_h="single",
            inside_v="single",
            border_size=8,
            border_color="FF0000",
        )

        # 验证所有边框属性设置正确
        for border_name in ("top", "bottom", "left", "right", "insideH", "insideV"):
            elem = self._get_border_elem(doc, "Table Grid", border_name)
            assert elem is not None, f"{border_name} 边框元素应存在"
            assert elem.get(qn("w:val")) == "single", f"{border_name} 样式应为 single"
            assert elem.get(qn("w:sz")) == "64", f"{border_name} 边框大小应为 8磅(64)"
            assert elem.get(qn("w:color")) == "FF0000", f"{border_name} 边框颜色应为 FF0000"

    def test_set_table_borders_partial(self) -> None:
        """测试部分设置表格边框（仅设置 top）."""
        doc = WordDocument()
        # 先创建一个新的自定义表格样式，确保干净的初始状态
        doc.format.create_table_style("测试边框样式")
        # 设置所有边框为 none 作为清理
        doc.format.set_table_borders(
            "测试边框样式",
            top="none",
            bottom="none",
            left="none",
            right="none",
            inside_h="none",
            inside_v="none",
        )
        # 现在只设置 top 边框
        doc.format.set_table_borders("测试边框样式", top="single", border_size=6, border_color="00FF00")

        # 验证 top 边框设置正确
        top_elem = self._get_border_elem(doc, "测试边框样式", "top")
        assert top_elem is not None, "top 边框元素应存在"
        assert top_elem.get(qn("w:val")) == "single"
        assert top_elem.get(qn("w:sz")) == "48"
        assert top_elem.get(qn("w:color")) == "00FF00"

        # 验证其他边框未被修改（仍为 none）
        for border_name in ("bottom", "left", "right", "insideH", "insideV"):
            elem = self._get_border_elem(doc, "测试边框样式", border_name)
            assert elem is not None, f"{border_name} 边框元素应存在"
            assert elem.get(qn("w:val")) == "none", f"{border_name} 应保持为 none，实际为 {elem.get(qn('w:val'))}"

    def test_set_table_borders_color_and_size_only(self) -> None:
        """测试仅设置边框颜色和大小（不改变边框样式）."""
        doc = WordDocument()
        # 创建样式并先设置所有边框
        doc.format.create_table_style("测试颜色样式")
        doc.format.set_table_borders(
            "测试颜色样式",
            top="single",
            bottom="single",
            left="single",
            right="single",
            inside_h="single",
            inside_v="single",
            border_size=8,
            border_color="000000",
        )

        # 验证初始状态 (border_size=8 → w:sz=64)
        top_elem = self._get_border_elem(doc, "测试颜色样式", "top")
        assert top_elem.get(qn("w:color")) == "000000"
        assert top_elem.get(qn("w:sz")) == "64"

        # 仅更新颜色和大小，不改变边框样式 (border_size=12 → w:sz=96)
        doc.format.set_table_borders("测试颜色样式", border_color="FF0000", border_size=12)

        # 验证颜色和大小已更新，但边框样式不变
        top_elem = self._get_border_elem(doc, "测试颜色样式", "top")
        assert top_elem.get(qn("w:val")) == "single", "边框样式应保持为 single"
        assert top_elem.get(qn("w:sz")) == "96", "边框大小应更新为 12磅(96)"
        assert top_elem.get(qn("w:color")) == "FF0000", "边框颜色应更新为 FF0000"

    def test_set_table_alignment(self) -> None:
        """测试设置表格对齐方式."""
        doc = WordDocument()
        doc.format.set_table_alignment("Table Grid", alignment="center")

    def test_set_table_paragraph(self) -> None:
        """测试设置表格段落格式."""
        doc = WordDocument()
        doc.format.set_table_paragraph(
            "Table Grid",
            alignment="center",
            space_before=0,
            space_after=6,
        )

    def test_table_style_all_format(self) -> None:
        """测试设置表格样式所有格式."""
        doc = WordDocument()
        doc.format.set_table_font("Table Grid", name="宋体", size=12, bold=True)
        doc.format.set_table_shading("Table Grid", fill="E6E6E6")
        doc.format.set_table_borders("Table Grid", top="single", bottom="single", border_size=8)
        doc.format.set_table_alignment("Table Grid", alignment="center")
        doc.format.set_table_paragraph("Table Grid", alignment="center")
