"""WordCell 模块测试."""

from __future__ import annotations

import tempfile
import os

from sword import WordDocument, WordTable
from docx.oxml.ns import qn


class TestWordCell:
    """WordCell 测试类."""

    def _get_cell_elem(self, doc: WordDocument, row: int, col: int):
        """获取单元格 XML 元素."""
        # 通过 document 内部结构获取单元格元素
        body = doc._inner.element.body
        tbl = body.find(qn("w:tbl"))
        if tbl is not None:
            tr = tbl.findall(qn("w:tr"))[row]
            tc = tr.findall(qn("w:tc"))[col]
            return tc
        return None

    def test_cell_text(self) -> None:
        """测试获取和设置单元格文本."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.text = "测试内容"
        assert cell.text == "测试内容"

    def test_cell_inner(self) -> None:
        """测试获取底层 Cell 对象."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        assert cell.inner is not None

    def test_cell_context_manager(self) -> None:
        """测试上下文管理器."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        with table.cell(0, 0) as cell:
            cell.text = "测试内容"
        assert table.cell(0, 0).text == "测试内容"

    def test_set_shading(self) -> None:
        """测试设置单元格底纹."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_shading("FFFF00")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        shd = tcPr.find(qn("w:shd"))
        assert shd is not None
        assert shd.get(qn("w:fill")) == "FFFF00"
        assert shd.get(qn("w:val")) == "clear"

    def test_set_borders(self) -> None:
        """测试设置单元格边框."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_borders(
            top="single",
            bottom="double",
            left="single",
            right="single",
            border_size=8,
            border_color="FF0000",
        )

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        tcBorders = tcPr.find(qn("w:tcBorders"))

        top = tcBorders.find(qn("w:top"))
        assert top.get(qn("w:val")) == "single"
        assert top.get(qn("w:sz")) == "64"
        assert top.get(qn("w:color")) == "FF0000"

        bottom = tcBorders.find(qn("w:bottom"))
        assert bottom.get(qn("w:val")) == "double"

    def test_set_margins(self) -> None:
        """测试设置单元格页边距."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_margins(top=10, bottom=20, left=30, right=40)

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        tcMar = tcPr.find(qn("w:tcMar"))

        top = tcMar.find(qn("w:top"))
        assert top.get(qn("w:w")) == "200"

        left = tcMar.find(qn("w:left"))
        assert left.get(qn("w:w")) == "600"

    def test_set_vertical_alignment(self) -> None:
        """测试设置单元格垂直对齐."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_vertical_alignment("center")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        vAlign = tcPr.find(qn("w:vAlign"))
        assert vAlign.get(qn("w:val")) == "center"

    def test_set_vertical_alignment_top(self) -> None:
        """测试设置单元格垂直对齐为顶部."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_vertical_alignment("top")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        vAlign = tcPr.find(qn("w:vAlign"))
        assert vAlign.get(qn("w:val")) == "top"

    def test_set_vertical_alignment_bottom(self) -> None:
        """测试设置单元格垂直对齐为底部."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_vertical_alignment("bottom")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        vAlign = tcPr.find(qn("w:vAlign"))
        assert vAlign.get(qn("w:val")) == "bottom"

    def test_set_horizontal_alignment_center(self) -> None:
        """测试设置单元格水平对齐为居中."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_horizontal_alignment("center")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        pPr = tcPr.find(qn("w:pPr"))
        jc = pPr.find(qn("w:jc"))
        assert jc.get(qn("w:val")) == "center"

    def test_set_horizontal_alignment_left(self) -> None:
        """测试设置单元格水平对齐为左对齐."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_horizontal_alignment("left")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        pPr = tcPr.find(qn("w:pPr"))
        jc = pPr.find(qn("w:jc"))
        assert jc.get(qn("w:val")) == "left"

    def test_set_horizontal_alignment_right(self) -> None:
        """测试设置单元格水平对齐为右对齐."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_horizontal_alignment("right")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        pPr = tcPr.find(qn("w:pPr"))
        jc = pPr.find(qn("w:jc"))
        assert jc.get(qn("w:val")) == "right"

    def test_set_horizontal_alignment_justify(self) -> None:
        """测试设置单元格水平对齐为两端对齐."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_horizontal_alignment("justify")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        pPr = tcPr.find(qn("w:pPr"))
        jc = pPr.find(qn("w:jc"))
        assert jc.get(qn("w:val")) == "both"

    def test_set_width_auto(self) -> None:
        """测试设置单元格宽度为自动."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_width(0, unit="auto")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        tcW = tcPr.find(qn("w:tcW"))
        assert tcW.get(qn("w:type")) == "auto"
        assert tcW.get(qn("w:w")) == "0"

    def test_set_width_dxa(self) -> None:
        """测试设置单元格宽度为固定值."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_width(1000, unit="dxa")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        tcW = tcPr.find(qn("w:tcW"))
        assert tcW.get(qn("w:type")) == "dxa"
        assert tcW.get(qn("w:w")) == "1000"

    def test_set_width_pct(self) -> None:
        """测试设置单元格宽度为百分比."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_width(50, unit="pct")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        tcW = tcPr.find(qn("w:tcW"))
        assert tcW.get(qn("w:type")) == "pct"
        assert tcW.get(qn("w:w")) == "50"

    def test_set_borders_partial(self) -> None:
        """测试部分设置单元格边框."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_borders(top="single", border_size=6, border_color="00FF00")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))
        tcBorders = tcPr.find(qn("w:tcBorders"))

        top = tcBorders.find(qn("w:top"))
        assert top.get(qn("w:val")) == "single"
        assert top.get(qn("w:sz")) == "48"
        assert top.get(qn("w:color")) == "00FF00"

    def test_all_format_methods(self) -> None:
        """测试单元格所有格式方法."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3)
        cell = table.cell(0, 0)
        cell.set_shading("E6E6E6")
        cell.set_borders(top="single", bottom="single", border_size=8)
        cell.set_margins(top=5, bottom=5, left=10, right=10)
        cell.set_vertical_alignment("center")
        cell.set_horizontal_alignment("center")
        cell.set_width(500, unit="dxa")

        tc = self._get_cell_elem(doc, 0, 0)
        tcPr = tc.find(qn("w:tcPr"))

        # 验证底纹
        shd = tcPr.find(qn("w:shd"))
        assert shd.get(qn("w:fill")) == "E6E6E6"

        # 验证边框
        tcBorders = tcPr.find(qn("w:tcBorders"))
        assert tcBorders.find(qn("w:top")).get(qn("w:val")) == "single"

        # 验证边距
        tcMar = tcPr.find(qn("w:tcMar"))
        assert tcMar.find(qn("w:top")).get(qn("w:w")) == "100"

        # 验证垂直对齐
        assert tcPr.find(qn("w:vAlign")).get(qn("w:val")) == "center"

        # 验证水平对齐
        pPr = tcPr.find(qn("w:pPr"))
        assert pPr.find(qn("w:jc")).get(qn("w:val")) == "center"

        # 验证宽度
        assert tcPr.find(qn("w:tcW")).get(qn("w:w")) == "500"

    def test_save_document_with_cell_format(
        self, tmp_path: tempfile.TemporaryDirectory
    ) -> None:
        """测试保存包含格式单元格的文档."""
        doc = WordDocument()
        table = WordTable(doc._inner, rows=3, cols=3, style="Table Grid")
        cell = table.cell(0, 0)
        cell.text = "标题"
        cell.set_shading("FFFF00")
        cell.set_horizontal_alignment("center")

        file_path = os.path.join(tmp_path, "cell_format_test.docx")
        doc.save(file_path)
        assert os.path.exists(file_path)
