"""Word 文档章节封装模块."""

from __future__ import annotations

from typing import Any

from docx.document import Document as DocxDocument

from sword.table import WordTable
from sword.paragraph import WordParagraph


class WordSection:
    """封装 Word 文档中的章节，支持添加标题、段落等内容."""

    def __init__(
        self,
        document: DocxDocument,
        title: str | None = None,
        number_counts: dict[int, int] | None = None,
        title_level: int = 1,
    ) -> None:
        """
        初始化章节.

        Args:
            document: 底层的 python-docx Document 对象.
            title: 章节标题（可选）.
            number_counts: 初始编号计数状态（可选），用于延续上一章节的编号.
            title_level: 标题级别（1-9，默认 1）.
        """
        self._doc = document
        self._title = title
        self._title_level = title_level
        self._number_counts = (
            number_counts
            if number_counts is not None
            else {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0, 7: 0, 8: 0, 9: 0}
        )
        if title:
            self.add_numbered_heading(title)

    def add_numbered_heading(self, text: str) -> None:
        """
        添加带自动编号的章节标题（层级编号如 1、1.1、1.1.1）.

        Args:
            text: 标题文本.
        """
        self._number_counts[self._title_level] += 1

        for l in range(self._title_level + 1, 10):
            self._number_counts[l] = 0

        number_parts = [
            str(self._number_counts[l]) for l in range(1, self._title_level + 1)
        ]
        number_str = ".".join(number_parts)
        full_text = f"{number_str} {text}"

        self._doc.add_heading(full_text, level=self._title_level)

    def add_paragraph(self, text: str = "", style: str | None = None) -> WordParagraph:
        """
        添加段落.

        Args:
            text: 段落文本.
            style: 段落样式（可选）.

        Returns:
            WordParagraph: 创建的段落封装对象。
        """
        para = self._doc.add_paragraph(text, style=style)
        return WordParagraph(para)

    def add_page_break(self) -> None:
        """添加分页符."""
        self._doc.add_page_break()

    def add_table(
        self,
        rows: int,
        cols: int,
        style: str | None = None,
    ) -> WordTable:
        """
        添加表格.

        Args:
            rows: 表格行数.
            cols: 表格列数.
            style: 表格样式（如 "Table Grid"，可选）.

        Returns:
            WordTable: 创建的表格封装对象.
        """
        return WordTable(self._doc, rows, cols, style)

    def add_section(self, title: str | None = None) -> WordSection:
        """
        添加子章节.

        Args:
            title: 章节标题（可选）.
            title_level: 标题级别（1-9，默认 1）.

        Returns:
            WordSection: 新创建的子章节对象.
        """
        return WordSection(self._doc, title, self._number_counts, self._title_level + 1)

    @property
    def title(self) -> str | None:
        """获取章节标题."""
        return self._title

    def get_number_counts(self) -> dict[int, int]:
        """获取当前编号计数状态（供外部保存延续）."""
        return self._number_counts.copy()

    def __enter__(self) -> WordSection:
        return self

    def __exit__(self, *args: Any) -> None:
        pass
